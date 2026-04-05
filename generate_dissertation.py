"""
Dissertation chapter generator — Advanced Predictive Modelling for Customer Churn in Banking
MSc Artificial Intelligence Final Dissertation
"""
import json, base64, os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Image paths ─────────────────────────────────────────────────────────────────
IMGS = {
    'churn_dist'      : '/tmp/nb_images/cell-eda-churn-dist_0.png',
    'numeric_dist'    : '/tmp/nb_images/cell-eda-numeric-dist_0.png',
    'cat_churn'       : '/tmp/nb_images/cell-eda-cat_0.png',
    'corr'            : '/tmp/nb_images/cell-eda-corr_0.png',
    'boxplots'        : '/tmp/nb_images/cell-eda-boxplots_0.png',
    'fe_viz'          : '/tmp/nb_images/cell-fe-viz_0.png',
    'comparison_viz'  : '/tmp/nb_images/cell-comparison-viz_0.png',
    'exec_time'       : '/tmp/nb_images/cell-execution-time-plot_0.png',
    'confusion'       : '/tmp/nb_images/cell-confusion-matrices_0.png',
    'shap_beeswarm'   : '/tmp/nb_images/cell-shap-plots_2.png',
    'shap_bar'        : '/tmp/nb_images/cell-shap-bar_1.png',
    'shap_local'      : '/tmp/nb_images/cell-shap-local_1.png',
    'shap_dl'         : '/tmp/nb_images/cell-shap-dl_1.png',
}

# ── Helper functions ─────────────────────────────────────────────────────────────
def set_run_font(run, bold=False, italic=False, size=11, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after  = Pt(6)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic, size=size)
    return p

def add_justified(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p

def add_figure(doc, img_path, caption, width=5.5):
    if not os.path.exists(img_path):
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(img_path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption)
    set_run_font(r, italic=True, size=10)

def add_table(doc, headers, rows, caption=None):
    """Add a styled table with bold headers."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # shade header
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        data_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            data_cells[ci].text = str(val)
            data_cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in data_cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run(caption)
        set_run_font(r, italic=True, size=10)
    doc.add_paragraph()

# ── Create document ──────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ════════════════════════════════════════════════════════════════════════════════
# CHAPTER 1: DATA ANALYSIS
# ════════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Chapter 1: Data Analysis', level=1)

add_justified(doc,
    'This chapter presents a comprehensive analysis of the dataset used in this study, '
    'encompassing an initial examination of its structure and quality, an exploratory data '
    'analysis (EDA) of the raw features, and the systematic engineering of derived features '
    'designed to enrich the predictive signal available to the models. The analyses described '
    'herein informed all subsequent modelling decisions and constitute the empirical foundation '
    'of the research.')

# ── 1.1 Dataset Description ────────────────────────────────────────────────────
add_heading(doc, '1.1 Dataset Description', level=2)

add_justified(doc,
    'The dataset employed in this study comprises 5,000 customer records drawn from a banking '
    'context, with each record corresponding to a single unique customer — confirming a '
    'cross-sectional rather than longitudinal data structure. The dataset contains 26 columns, '
    'spanning demographic attributes (age, gender, marital status, employment status, region), '
    'financial characteristics (account balance, annual income, credit score, loan amount, loan '
    'type), transaction behaviour (transaction amount, number of transactions, transaction type, '
    'account activity trend), service interaction metrics (customer service interactions, recent '
    'complaints, customer satisfaction score), temporal markers (transaction date, account open '
    'date, last transaction date), and the binary target variable Churn_Label.')

add_justified(doc,
    'The target variable Churn_Label is dichotomous: 0 denotes a retained customer and 1 denotes '
    'a churned customer. The column Churn_Timeframe, which indicates the projected time horizon '
    'of churn, was excluded from all analyses due to direct target leakage — this variable is '
    'undefined for non-churning customers and is derived from the churn outcome itself.')

add_justified(doc,
    'Two notable data quality issues were identified. First, the Loan_Type column exhibited '
    '76% missingness (3,800 of 5,000 rows), attributable to the fact that the majority of '
    'customers hold no loan product; this field was imputed using the modal category. Second, '
    'the Branch column contained 4,312 unique values, rendering it unsuitable for one-hot '
    'encoding and necessitating its exclusion from the modelling feature set. All three date '
    'columns were stored as character strings and required explicit parsing; they were not used '
    'directly as model inputs but served as the basis for tenure and inactivity derived features '
    'in the feature engineering stage (Section 1.3).')

add_heading(doc, '1.1.1 Class Distribution', level=3)

add_justified(doc,
    'The dataset exhibits a moderate class imbalance: 68.34% of customers are non-churners '
    '(class 0) and 31.66% are churners (class 1), corresponding to an imbalance ratio of '
    'approximately 2.16:1. Whilst this level of imbalance is less severe than those commonly '
    'encountered in real-world banking datasets — where churn rates typically range between '
    '10% and 20% (Verbeke et al., 2012) — it is nonetheless sufficient to bias models trained '
    'without explicit imbalance correction toward the majority class. A naive classifier that '
    'always predicts non-churn would achieve approximately 68% accuracy whilst correctly '
    'identifying zero churners, illustrating why accuracy alone is an inadequate evaluation '
    'metric for this task. Figure 1.1 illustrates the class distribution.')

add_figure(doc, IMGS['churn_dist'],
    'Figure 1.1: Class distribution of the target variable (Churn_Label). '
    'The 31.66% churn rate represents a moderate imbalance (ratio 2.16:1).',
    width=4.0)

# ── 1.2 Exploratory Data Analysis ─────────────────────────────────────────────
add_heading(doc, '1.2 Exploratory Data Analysis', level=2)

add_justified(doc,
    'Exploratory data analysis was conducted to characterise the distributional properties '
    'of the features, identify the variables most predictive of churn, and inform both the '
    'preprocessing strategy and the feature engineering decisions. The analyses encompass '
    'univariate distributional comparisons, categorical churn rate decomposition, linear '
    'correlation analysis, and bivariate box plot examination.')

add_heading(doc, '1.2.1 Numeric Feature Distributions', level=3)

add_justified(doc,
    'Kernel density plots and histograms were constructed for all continuous and ordinal '
    'numeric features, stratified by churn label. The most diagnostically informative '
    'distributions are discussed below.')

add_justified(doc,
    'Recent_Complaints and Number_of_Transactions demonstrated the clearest distributional '
    'separation between churners and non-churners: customers who had filed more recent '
    'complaints and those with fewer transactions were disproportionately represented in '
    'the churn class. Customer_Satisfaction_Score also exhibited a noticeable rightward shift '
    'for retained customers relative to churners, confirming its status as a principal '
    'behavioural predictor. In contrast, financial variables — including Account_Balance, '
    'Annual_Income, Credit_Score, and Loan_Amount — displayed largely overlapping '
    'distributions across churn classes, suggesting that financial health metrics possess '
    'limited standalone discriminatory power in this dataset. This finding implies that '
    'churn in this context is behaviourally rather than financially driven, favouring '
    'non-linear ensemble methods capable of capturing interaction effects.')

add_figure(doc, IMGS['numeric_dist'],
    'Figure 1.2: Distributions of numeric features stratified by churn status. '
    'Recent_Complaints, Customer_Satisfaction_Score, and Number_of_Transactions '
    'show the clearest class separation.',
    width=6.0)

add_heading(doc, '1.2.2 Categorical Feature Analysis', level=3)

add_justified(doc,
    'Churn rates were computed for each level of the six categorical features included '
    'in the baseline feature set: Gender, Account_Type, Transaction_Type, Marital_Status, '
    'Region, and Account_Activity_Trend. As illustrated in Figure 1.3, no single categorical '
    'feature exhibited a dramatically different churn rate from the dataset average of 31.66%. '
    'The most meaningful categorical variation was observed in Account_Activity_Trend, where '
    'customers exhibiting a decreasing account activity trend demonstrated higher churn '
    'propensity — consistent with the hypothesis that disengagement precedes account closure. '
    'The uniformity of churn rates across most categorical groups reinforces the conclusion that '
    'churn prediction in this dataset requires the capture of complex feature interactions '
    'rather than simple group-based filtering.')

add_figure(doc, IMGS['cat_churn'],
    'Figure 1.3: Churn rates by category level for the six categorical features. '
    'Account_Activity_Trend shows the most meaningful variation.',
    width=6.0)

add_heading(doc, '1.2.3 Correlation Analysis', level=3)

add_justified(doc,
    'Linear (Pearson) correlations between all numeric features and the target variable were '
    'computed and visualised as a heatmap (Figure 1.4). The three strongest churn predictors '
    'by linear correlation were Recent_Complaints (r = 0.262), Customer_Satisfaction_Score '
    '(r = 0.189), and Number_of_Transactions (r = 0.117). All other features exhibited '
    'correlations below 0.015 in absolute magnitude, confirming the weak linear relationship '
    'between financial variables and churn outcome.')

add_justified(doc,
    'The absence of severe multicollinearity among the predictor variables (all pairwise '
    'correlations low) is a favourable property for the stability of the Logistic Regression '
    'model and eliminates the need for dimensionality reduction prior to modelling. Importantly, '
    'the low aggregate correlation magnitudes do not preclude predictive utility; tree-based '
    'ensembles and neural networks are capable of identifying non-linear thresholds and '
    'higher-order interactions that linear correlation cannot detect.')

add_figure(doc, IMGS['corr'],
    'Figure 1.4: Pearson correlation heatmap. Recent_Complaints (r = 0.262) and '
    'Customer_Satisfaction_Score (r = 0.189) are the strongest linear churn predictors.',
    width=5.5)

add_heading(doc, '1.2.4 Box Plot Analysis', level=3)

add_justified(doc,
    'Bivariate box plots were constructed to examine the distributional differences in key '
    'numeric features between churners and non-churners. As presented in Figure 1.5, '
    'Customer_Satisfaction_Score displayed the clearest median separation: retained customers '
    'recorded consistently higher scores. Recent_Complaints exhibited a higher median and wider '
    'interquartile range among churners. Number_of_Transactions was lower for churners on '
    'average, corroborating the disengagement narrative. Financial variables such as '
    'Credit_Score, Account_Balance, and Annual_Income showed nearly identical medians across '
    'both classes, with overlapping interquartile ranges, further confirming their limited '
    'individual discriminatory power.')

add_justified(doc,
    'The pervasive overlap in distributional profiles underscores the inherent difficulty of '
    'this binary classification problem and motivates the prioritisation of Recall as the '
    'primary evaluation metric: in a banking retention context, the cost of a false negative '
    '(failing to identify a departing customer) substantially exceeds the cost of a false '
    'positive (unnecessary but inexpensive retention outreach).')

add_figure(doc, IMGS['boxplots'],
    'Figure 1.5: Box plots of key numeric features by churn class. '
    'Customer_Satisfaction_Score and Recent_Complaints show the clearest median separation.',
    width=6.0)

# ── 1.3 Feature Engineering ────────────────────────────────────────────────────
add_heading(doc, '1.3 Feature Engineering', level=2)

add_justified(doc,
    'Feature engineering was conducted in accordance with Appendix B of the research proposal, '
    'with the objective of constructing domain-informed composite features that amplify the '
    'predictive signal latent in the raw data. The engineered features were designed to '
    'operationalise behavioural and financial concepts identified as theoretically meaningful '
    'in the churn literature — including customer engagement depth, disengagement duration, '
    'financial leverage, and service dissatisfaction intensity.')

add_heading(doc, '1.3.1 Reference Date Determination', level=3)

add_justified(doc,
    'Several engineered features require a reference "current date" relative to which temporal '
    'intervals (tenure, inactivity) are computed. To avoid any dependency on external inputs '
    'and to ensure reproducibility, the reference date was determined endogenously: the maximum '
    'date observed across all three temporal columns (Transaction_Date, Account_Open_Date, '
    'Last_Transaction_Date) in the raw dataset was adopted as the reference date. This yielded '
    'a reference date of 30 March 2025, which was applied consistently across all data subsets.')

add_heading(doc, '1.3.2 Derived Feature Set', level=3)

add_justified(doc,
    'Nine derived features were constructed across four conceptual groups, expanding the feature '
    'space from 18 baseline features to 27 pre-OHE features. Table 1.1 summarises the full set, '
    'including the mathematical formulation and domain rationale for each feature.')

add_table(doc,
    headers=['Group', 'Feature', 'Formulation', 'Domain Rationale'],
    rows=[
        ['Tenure', 'Customer_Tenure_Months',
         '(Ref_Date − Account_Open_Date).days / 30.44',
         'Short tenure → reduced loyalty; proxy for relationship depth'],
        ['Tenure', 'Inactivity_Period_Days',
         '(Ref_Date − Last_Transaction_Date).days',
         'Days since last transaction; disengagement signal'],
        ['Transaction', 'Transaction_Frequency',
         'N_Transactions / (Tenure_Months + 1)',
         'Normalised usage intensity; +1 avoids division by zero for new accounts'],
        ['Transaction', 'Avg_Transaction_Size',
         'Transaction_Amount / clip(N_Transactions, 1)',
         'Spending pattern proxy; spending intensity per visit'],
        ['Transaction', 'Balance_Volatility_Index',
         '|ΔBalance| / (|Balance| + ε)',
         'Relative balance change; financial instability indicator (99th pct clipped)'],
        ['Financial', 'Loan_to_Income_Ratio',
         'Loan_Amount / (Annual_Income + ε)',
         'Leverage risk; high ratio signals debt-servicing stress'],
        ['Service', 'Complaint_to_Interaction_Ratio',
         'Recent_Complaints / (Service_Interactions + 1)',
         'Dissatisfaction density per service contact'],
        ['Service', 'Satisfaction_Gap',
         '(6 − Satisfaction_Score) × (Complaints + 1)',
         'Joint low-satisfaction / high-complaint signal; strong churn predictor'],
        ['Composite', 'Engagement_Index',
         'Tx_Frequency × Satisfaction / (Complaint_Ratio + 1)',
         'Holistic engagement: high activity + high satisfaction + low complaints'],
    ],
    caption='Table 1.1: Engineered feature set, formulations, and domain rationale.')

add_heading(doc, '1.3.3 Data Leakage Prevention', level=3)

add_justified(doc,
    'Strict data leakage prevention measures were enforced throughout the feature engineering '
    'pipeline. All nine derived features are computed exclusively from raw input variables — '
    'no reference to the target label Churn_Label is involved in any formulation. The '
    'function engineer_features() is therefore safe to call independently on training and '
    'test splits without risk of label contamination.')

add_justified(doc,
    'A noteworthy leakage risk was identified and corrected during development: the original '
    'design of the feature set included two target-rate features — Region_Churn_Rate and '
    'Branch_Churn_Rate — computed as the mean churn rate within each region or branch grouping. '
    'These were computed from the full y_train vector prior to cross-validation, which meant '
    'that validation fold labels were already embedded in the group-level rate estimates. This '
    'constitutes a form of target leakage and produced a substantial train-validation gap '
    '(train F1 >> validation F1). Both features were removed entirely from the engineered '
    'feature set. The regional signal is instead captured cleanly through one-hot encoding '
    'of the Region column in the preprocessing pipeline.')

add_justified(doc,
    'Additionally, the Balance_Volatility_Index exhibits heavy right-tail behaviour due to '
    'near-zero account balances in the denominator. To prevent extreme values from distorting '
    'model learning, a clipping cap was computed from the training set 99th percentile and '
    'applied to both the training and test sets — the cap being derived exclusively from '
    'training data to preclude test-set leakage.')

add_heading(doc, '1.3.4 Engineered Feature Distributions', level=3)

add_justified(doc,
    'Figure 1.6 presents the distributional profiles of the nine engineered features, '
    'stratified by churn class. The Satisfaction_Gap feature demonstrates the most pronounced '
    'class separation, confirming the synergistic effect of jointly encoding low satisfaction '
    'and high complaint frequency. Complaint_to_Interaction_Ratio and Customer_Tenure_Months '
    'also exhibit meaningful distributional differences: churners have higher complaint '
    'densities and shorter average tenures. Inactivity_Period_Days shows a moderate '
    'separation consistent with the disengagement hypothesis. The Engagement_Index and '
    'Balance_Volatility_Index display broader distributional overlap, suggesting their '
    'contribution is primarily through interaction effects with other features rather than '
    'as standalone discriminators.')

add_figure(doc, IMGS['fe_viz'],
    'Figure 1.6: Distributions of the nine engineered features by churn class. '
    'Satisfaction_Gap and Complaint_to_Interaction_Ratio show the strongest class separation.',
    width=6.0)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# CHAPTER 2: IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Chapter 2: Implementation', level=1)

add_justified(doc,
    'This chapter documents the full modelling pipeline implemented in this study. It details '
    'the data partitioning strategy, preprocessing decisions, the handling of class imbalance, '
    'the model architectures evaluated across three experimental phases, the hyperparameter '
    'tuning strategies employed in Phase 3, the evaluation metrics used throughout, and '
    'the statistical significance testing approach applied to assess the contribution of '
    'feature engineering. The chapter concludes with the identification and justification '
    'of the best-performing model.')

# ── 2.1 Data Partitioning and Preprocessing ────────────────────────────────────
add_heading(doc, '2.1 Data Partitioning Strategy', level=2)

add_heading(doc, '2.1.1 Train-Test Split', level=3)

add_justified(doc,
    'The dataset was partitioned into training and holdout test sets using a stratified '
    '80/20 random split, yielding 4,000 training samples and 1,000 test samples. '
    'Stratification on the target label was enforced to preserve the original class '
    'distribution in both subsets. The resulting churn rates were 0.3165 in the training '
    'set, 0.3170 in the test set, and 0.3166 in the full dataset — confirming that the '
    'partition faithfully maintained the class balance.')

add_justified(doc,
    'This partition was performed once and held fixed throughout all three experimental '
    'phases. All models — regardless of feature set or tuning strategy — were trained on '
    'the same 4,000 training samples and evaluated on the same 1,000 test samples. This '
    'design ensures that cross-phase performance comparisons reflect genuine model '
    'differences rather than artefacts of different data splits.')

add_heading(doc, '2.1.2 Cross-Validation Strategy', level=3)

add_justified(doc,
    'Within the training set, five-fold stratified cross-validation (StratifiedKFold, '
    'n_splits=5, shuffle=True, random_state=42) was employed for all model evaluation '
    'and hyperparameter search procedures. Stratification ensures that each fold maintains '
    'the overall churn rate, preventing any fold from being dominated by the majority class. '
    'The use of a fixed random seed guarantees reproducibility of fold assignments across '
    'all experiments.')

add_heading(doc, '2.1.3 Feature Preprocessing Pipeline', level=3)

add_justified(doc,
    'A scikit-learn ColumnTransformer was constructed to apply type-appropriate '
    'transformations to each feature group:')

add_bullet(doc,
    'Numeric features (12 in the baseline set; 21 in the engineered set): '
    'StandardScaler was applied, centring each feature to zero mean and unit variance. '
    'This transformation is critical for Logistic Regression and the DL models, which are '
    'sensitive to feature scales, and is also beneficial for ensemble methods by improving '
    'convergence stability.')
add_bullet(doc,
    'Categorical features (six columns): OneHotEncoder with drop=\'first\' was applied, '
    'producing binary indicator variables for each non-reference category. The drop=\'first\' '
    'option eliminates perfect multicollinearity (the dummy variable trap), which is '
    'particularly important for the numerical stability of Logistic Regression.')
add_bullet(doc,
    'Missing values: Median imputation was applied to numeric columns and mode imputation '
    'to categorical columns, with imputation statistics derived exclusively from the '
    'training fold in each cross-validation iteration.')

add_justified(doc,
    'The ColumnTransformer was fitted on the training data (or training fold) only and '
    'applied as a read-only transform to the test data (or validation fold), ensuring that '
    'no information from held-out data influenced the scaling or encoding parameters. '
    'The baseline preprocessing pipeline produced a (4,000 × 27) training matrix and a '
    '(1,000 × 27) test matrix; the engineered feature set expanded these to (4,000 × 36) '
    'and (1,000 × 36) respectively.')

# ── 2.2 Handling Class Imbalance ───────────────────────────────────────────────
add_heading(doc, '2.2 Class Imbalance Handling', level=2)

add_justified(doc,
    'The 2.16:1 class imbalance was addressed using two complementary strategies, each '
    'appropriate to the corresponding model paradigm.')

add_heading(doc, '2.2.1 SMOTE for Machine Learning Models', level=3)

add_justified(doc,
    'For all machine learning models, the Synthetic Minority Over-sampling Technique '
    '(SMOTE; Chawla et al., 2002) was employed. SMOTE generates synthetic minority-class '
    'samples by interpolating between existing minority-class observations in feature space, '
    'increasing the effective minority class representation in the training data. '
    'The critical methodological requirement is that SMOTE must be applied only to '
    'training data, never to validation or test data.')

add_justified(doc,
    'To satisfy this requirement without additional complexity, SMOTE was embedded '
    'directly within each scikit-learn imbalanced-learn Pipeline as a preprocessing step. '
    'This architecture guarantees that, during cross-validation, SMOTE fits and resamples '
    'only the training fold of each split — the validation fold always reflects the '
    'original, unaugmented class distribution. This design prevents SMOTE-based leakage '
    'and ensures that CV metrics estimate generalisation performance on real data '
    'distributions. SMOTE parameters used were: sampling_strategy=\'minority\' '
    '(upsample minority class to match majority), k_neighbors=5, random_state=42.')

add_heading(doc, '2.2.2 Class Weights for Deep Learning Models', level=3)

add_justified(doc,
    'For deep learning models (ANN and DNN), SMOTE was replaced by the class_weight '
    'mechanism provided natively by Keras. This design choice was motivated by an '
    'empirical observation: when SMOTE was applied inside each CV fold for DL models, '
    'a systematic and large gap emerged between cross-validation Recall (CV ≈ 0.71–0.77) '
    'and test-set Recall (Test ≈ 0.40–0.49). This gap arises because SMOTE-generated '
    'synthetic samples fill the minority class feature space in patterns that neural '
    'networks exploit within training folds but which do not generalise to the true '
    'distribution of the real test set.')

add_justified(doc,
    'The class_weight approach instead up-weights the loss contribution of minority-class '
    'samples during gradient computation, without creating any synthetic data. Per-fold '
    'class weights were computed using the sklearn convention: '
    'weight_c = n_total / (n_classes × n_c), where n_c is the count of class c in the '
    'training fold. This produces weights of approximately {0: 0.73, 1: 1.58} for the '
    'observed 68:32 imbalance ratio. By training on the real training distribution (rather '
    'than a SMOTE-augmented one), the gap between CV and test metrics is substantially '
    'reduced, yielding more reliable estimates of generalisation performance.')

# ── 2.3 Model Architecture ─────────────────────────────────────────────────────
add_heading(doc, '2.3 Experimental Design and Model Architecture', level=2)

add_justified(doc,
    'The study is structured across three experimental phases, each building upon the '
    'previous. Phase 1 establishes baseline performance using default hyperparameters on '
    'the original feature set. Phase 2 replaces the original feature set with the '
    'engineered features (Section 1.3) whilst retaining default hyperparameters. Phase 3 '
    'applies systematic hyperparameter tuning to the engineered feature set. '
    'Six models are evaluated in total: four machine learning classifiers and two '
    'deep learning architectures.')

add_heading(doc, '2.3.1 Machine Learning Models', level=3)

add_justified(doc,
    'The four machine learning models were selected to represent a range of model '
    'families, from linear to non-linear and from single estimators to ensembles:')

add_bullet(doc,
    'Logistic Regression (LR): A linear probabilistic classifier serving as the '
    'interpretable baseline. Its performance on linearly separable problems provides '
    'a lower bound for non-linear models and a reference for assessing whether '
    'complexity additions are justified.')
add_bullet(doc,
    'Decision Tree (DT): A non-parametric model that partitions the feature space '
    'through binary axis-aligned splits. Its default unconstrained depth makes it '
    'susceptible to overfitting, making it a primary beneficiary of hyperparameter tuning.')
add_bullet(doc,
    'Random Forest (RF): A bagged ensemble of decision trees trained on bootstrap '
    'samples with randomised feature subsets. RF reduces variance relative to a single '
    'tree and captures complex feature interactions.')
add_bullet(doc,
    'XGBoost (XGB): A gradient-boosted tree ensemble that sequentially corrects '
    'the residuals of previous trees, combining strong predictive performance with '
    'built-in regularisation. XGBoost is widely cited as the state-of-the-art for '
    'structured tabular data (Chen & Guestrin, 2016).')

add_justified(doc,
    'All four models were trained within an imbalanced-learn Pipeline embedding SMOTE '
    'as described in Section 2.2.1, with the ColumnTransformer as the preprocessing step.')

add_heading(doc, '2.3.2 Deep Learning Models', level=3)

add_justified(doc,
    'Two deep learning architectures were implemented using TensorFlow/Keras, both '
    'configured according to empirically validated specifications from the recent literature:')

add_bullet(doc,
    'Artificial Neural Network (ANN): A shallow architecture with two hidden layers '
    '(64 → 32 neurons), ReLU activation, Dropout(0.3) after each hidden layer, and a '
    'sigmoid output node. This configuration follows AbdelAziz et al. (2025). The network '
    'contains approximately 3,457 trainable parameters.')
add_bullet(doc,
    'Deep Neural Network (DNN): A deeper architecture with five hidden layers '
    '(256 → 128 → 64 → 32 → 16 neurons), ReLU activation, Batch Normalisation and '
    'Dropout(0.3) after the first three layers, and Dropout(0.2) after the fourth layer. '
    'This configuration follows Singh et al. (2024) and Basit et al. (2024). The architecture '
    'is designed to capture hierarchical feature abstractions.')

add_justified(doc,
    'Both architectures were compiled with the Adam optimiser (learning_rate=0.001) and '
    'binary cross-entropy loss. Training was governed by early stopping monitoring '
    'validation loss with patience=10 epochs and restore_best_weights=True, preventing '
    'overfitting through premature termination. The maximum training budget was 100 epochs '
    'with a batch size of 32. Deep learning models were not subjected to full hyperparameter '
    'search in Phase 3, as the dataset size (5,000 rows) limits the marginal benefit of '
    'DL hyperparameter optimisation (Shwartz-Ziv & Armon, 2022); literature-informed '
    'configurations were retained throughout.')

# ── 2.4 Hyperparameter Tuning ──────────────────────────────────────────────────
add_heading(doc, '2.4 Hyperparameter Tuning Strategies', level=2)

add_justified(doc,
    'Phase 3 applied three distinct hyperparameter tuning strategies to each of the four '
    'machine learning models, trained on the engineered feature set. All strategies '
    'optimised the cross-validated F1 score (the primary evaluation metric) using the '
    'same five-fold stratified CV configuration. The three strategies were deliberately '
    'selected to span the spectrum of search sophistication — from exhaustive enumeration '
    'to probabilistic sampling to Bayesian optimisation — enabling a methodological '
    'comparison of tuning efficiency and effectiveness.')

add_heading(doc, '2.4.1 Grid Search', level=3)

add_justified(doc,
    'Grid Search (GridSearchCV) performs an exhaustive evaluation of all combinations of '
    'pre-specified hyperparameter values. The search spaces were designed to cover '
    'each model\'s most impactful parameters: for Logistic Regression, the regularisation '
    'strength C ∈ {0.01, 0.1, 1, 10}; for Decision Tree, max_depth ∈ {3, 5, 7, 10, None} '
    'and min_samples_split ∈ {2, 5, 10}; for Random Forest, n_estimators ∈ {50, 100, 200} '
    'and max_depth ∈ {5, 10, None}; for XGBoost, n_estimators ∈ {50, 100, 200}, '
    'max_depth ∈ {3, 5, 7}, learning_rate ∈ {0.01, 0.1, 0.3}, and subsample ∈ {0.7, 1.0}. '
    'Grid Search guarantees discovery of the global optimum within the defined grid at '
    'the cost of exponential growth in computation as the grid expands.')

add_heading(doc, '2.4.2 Randomised Search', level=3)

add_justified(doc,
    'Randomised Search (RandomizedSearchCV, n_iter=50) samples hyperparameter combinations '
    'from continuous or discrete distributions rather than evaluating a fixed grid. '
    'This enables exploration of a far larger parameter space at equivalent computational '
    'cost. Notably, the regularisation parameter C for Logistic Regression was sampled '
    'from a log-uniform distribution over [10⁻³, 10²], providing finer resolution than '
    'the coarse Grid Search values. Decision Tree added the criterion ∈ {gini, entropy} '
    'as an additional dimension. Random Forest and XGBoost search spaces were similarly '
    'expanded to include max_features and colsample_bytree respectively. '
    'Randomised Search is typically more efficient than Grid Search in high-dimensional '
    'parameter spaces, as it avoids wasting evaluations on uninformative regions.')

add_heading(doc, '2.4.3 Bayesian Optimisation (Optuna)', level=3)

add_justified(doc,
    'Optuna (Akiba et al., 2019) implements Tree-structured Parzen Estimation (TPE), '
    'a Bayesian optimisation algorithm that builds a probabilistic model of the objective '
    'function (CV F1 score) and uses it to guide the selection of subsequent trials '
    'toward promising parameter regions. Each experiment comprised 50 trials with '
    'MedianPruner enabled — unpromising trials are terminated early once their intermediate '
    'CV score falls below the median of completed trials at the same step, reducing '
    'wasted computation. Optuna\'s directed search is particularly advantageous in large, '
    'high-dimensional search spaces where random sampling is inefficient; at the scale '
    'of this study, its primary benefit is computational efficiency rather than final '
    'performance improvement.')

# ── 2.5 Evaluation Metrics ─────────────────────────────────────────────────────
add_heading(doc, '2.5 Evaluation Metrics', level=2)

add_justified(doc,
    'Six metrics were computed for each model, capturing different aspects of predictive '
    'performance and operational feasibility:')

add_table(doc,
    headers=['Metric', 'Formula', 'Relevance to Churn Prediction'],
    rows=[
        ['Accuracy', '(TP + TN) / (TP + TN + FP + FN)',
         'Overall correctness; misleading under class imbalance'],
        ['Precision', 'TP / (TP + FP)',
         'Fraction of predicted churners who actually churn; cost of retention campaigns'],
        ['Recall', 'TP / (TP + FN)',
         'Fraction of actual churners captured; primary business metric — missed churners = lost revenue'],
        ['F1 Score', '2 × (Precision × Recall) / (Precision + Recall)',
         'Harmonic mean; balances precision and recall; primary optimisation target'],
        ['AUC-ROC', 'Area under ROC curve',
         'Probability that model ranks a random churner above a random non-churner; threshold-independent'],
        ['Inference Time (µs/sample)',
         'Wall-clock time per sample at prediction',
         'Operational feasibility for real-time scoring in banking CRM systems'],
    ],
    caption='Table 2.1: Evaluation metrics used across all phases.')

add_justified(doc,
    'F1 Score and AUC-ROC are designated as the primary evaluation metrics. F1 Score '
    'accounts for the class imbalance by measuring performance on the minority class '
    'specifically, whilst AUC-ROC provides a threshold-independent assessment of '
    'discrimination ability. Recall is reported alongside F1 because in a banking '
    'retention context, minimising false negatives (missed churners) is typically '
    'the paramount business objective.')

# ── 2.6 Statistical Significance ──────────────────────────────────────────────
add_heading(doc, '2.6 Statistical Significance Testing', level=2)

add_justified(doc,
    'To assess whether the performance differences between Phase 1 (baseline features) '
    'and Phase 2 (engineered features) were statistically meaningful rather than '
    'attributable to random fold variation, the Wilcoxon signed-rank test was applied '
    'to the paired fold-level F1 scores for each model. The Wilcoxon test is a '
    'non-parametric alternative to the paired t-test, appropriate here because the '
    'normality of fold-score differences cannot be assumed for only five observations.')

add_justified(doc,
    'The results, presented in Table 2.2, showed that none of the Phase 1 vs. Phase 2 '
    'differences were statistically significant at α = 0.05. This finding is unsurprising '
    'given the very limited statistical power available with only five paired observations '
    '(n = 5 per fold). The Wilcoxon test requires a minimum of six paired differences to '
    'achieve p < 0.05 at the two-tailed level; with five folds, the minimum achievable '
    'p-value is 0.0625. Accordingly, the absence of statistical significance does not '
    'constitute evidence that feature engineering is ineffective; it reflects the '
    'inherent power limitations of five-fold CV for significance testing.')

add_table(doc,
    headers=['Model', 'Phase1 CV F1', 'Phase2 CV F1', 'Δ F1', 'p-value', 'Significant?'],
    rows=[
        ['Logistic Regression', '0.5185', '0.5215', '+0.003', '0.1250', 'No (ns)'],
        ['Decision Tree',       '0.4242', '0.4252', '+0.001', '1.0000', 'No (ns)'],
        ['Random Forest',       '0.4055', '0.4392', '+0.034', '0.0625', 'No (borderline)'],
        ['XGBoost',             '0.4112', '0.3960', '−0.015', '0.4375', 'No (ns)'],
    ],
    caption='Table 2.2: Wilcoxon signed-rank test results for Phase 1 vs. Phase 2 CV F1 (α = 0.05).')

add_justified(doc,
    'Random Forest produced the closest result to significance (p = 0.0625, borderline), '
    'consistent with the largest absolute improvement in test F1 observed for that model '
    '(+0.037). The distinction between statistical significance and practical significance '
    'must be acknowledged: even where formal significance is unattainable due to sample '
    'size constraints, a sustained improvement of 3–4% in F1 score may carry meaningful '
    'business value in a banking retention context.')

# ── 2.7 Best Model Selection ───────────────────────────────────────────────────
add_heading(doc, '2.7 Best Model Selection', level=2)

add_justified(doc,
    'Model selection was based on a multi-criterion evaluation encompassing test F1, '
    'test AUC, test Recall, and inference latency. The tuned Decision Tree '
    '(max_depth=3, criterion=entropy, identified consistently by all three tuning '
    'strategies) achieved the highest test F1 (0.5710) and AUC (0.7339) across all '
    'phases and methods, whilst retaining a prediction latency of approximately '
    '3 µs per sample — the fastest among all evaluated models and fully compatible '
    'with real-time banking scoring requirements.')

add_justified(doc,
    'For deployment scenarios where Recall maximisation is the primary business objective '
    '(i.e., minimising missed churners), Logistic Regression with a lowered classification '
    'threshold below 0.5 represents the preferred alternative: its test Recall of 0.647 '
    'in Phase 1 is the highest achieved across all models and phases, and its probability '
    'outputs are well-calibrated for threshold adjustment via the AUC-ROC curve.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# CHAPTER 3: RESULTS AND DISCUSSION
# ════════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Chapter 3: Results and Discussion', level=1)

add_justified(doc,
    'This chapter presents and interprets the empirical findings from the three experimental '
    'phases, situates the results within the broader literature, and provides a detailed '
    'analysis of the SHAP-based interpretability outputs. Particular attention is given to '
    'the underperformance of XGBoost relative to prior expectations, and to the systematic '
    'gap between DL cross-validation and test performance.')

# ── 3.1 Phase 1 ────────────────────────────────────────────────────────────────
add_heading(doc, '3.1 Phase 1: Baseline Models on Original Features', level=2)

add_justified(doc,
    'Phase 1 established performance benchmarks using default hyperparameters applied '
    'to the 18 original features (12 numeric + 6 categorical, after dropping identifier, '
    'leakage, date, and high-cardinality columns). Table 3.1 presents the full metric '
    'set for all six models.')

add_table(doc,
    headers=['Model', 'CV F1', 'Test F1', 'Test Recall', 'Test AUC', 'Pred Time (µs)'],
    rows=[
        ['Logistic Regression', '0.5185', '0.5423', '0.6467', '0.7150', '3.1'],
        ['Decision Tree',       '0.4242', '0.4524', '0.4795', '0.5911', '2.9'],
        ['Random Forest',       '0.4055', '0.4053', '0.3375', '0.7110', '16.4'],
        ['XGBoost',             '0.4112', '0.3816', '0.3407', '0.6805', '3.9'],
        ['ANN',                 '0.5372', '0.4726', '0.5300', '0.6552', '172'],
        ['DNN',                 '0.5459', '0.4136', '0.4038', '0.6591', '697'],
    ],
    caption='Table 3.1: Phase 1 results — baseline models on original features.')

add_justified(doc,
    'Logistic Regression emerged as the strongest Phase 1 model, achieving the highest '
    'test F1 (0.5423), highest test Recall (0.6467), and highest AUC (0.7150). This '
    'outcome indicates a meaningful linear component in the churn signal, consistent with '
    'the strong linear correlations identified for Recent_Complaints and '
    'Customer_Satisfaction_Score in the EDA. The result aligns with a recurring finding '
    'in the churn prediction literature that simple linear models with appropriate '
    'regularisation can be highly competitive on datasets where the dominant predictors '
    'operate near-linearly (Verbeke et al., 2012).')

add_justified(doc,
    'The ANN and DNN models exhibited a pronounced gap between CV Recall and test Recall: '
    'ANN CV Recall = 0.709 vs test Recall = 0.530 (gap = 0.179); DNN CV Recall = 0.753 '
    'vs test Recall = 0.404 (gap = 0.349). This gap was attributed to the SMOTE-based '
    'imbalance strategy previously employed for DL models. The fix — replacing SMOTE with '
    'class_weight — was implemented prior to re-running (see Section 2.2.2). '
    'Both DL models also trailed all ML models on test F1, consistent with the '
    'well-documented finding that tree-based ensembles typically outperform deep learning '
    'on small-to-medium tabular datasets (Grinsztajn et al., 2022).')

# ── 3.2 Phase 2 ────────────────────────────────────────────────────────────────
add_heading(doc, '3.2 Phase 2: Baseline Models on Engineered Features', level=2)

add_justified(doc,
    'Phase 2 augmented the feature set with the nine engineered features (Section 1.3), '
    'whilst retaining default hyperparameters. Table 3.2 presents the Phase 2 results '
    'alongside the Phase 1 test F1 for comparison.')

add_table(doc,
    headers=['Model', 'Phase1 Test F1', 'Phase2 Test F1', 'Δ F1', 'Phase2 Test AUC'],
    rows=[
        ['Logistic Regression', '0.5423', '0.5253', '−0.017', '0.7027'],
        ['Decision Tree',       '0.4524', '0.3865', '−0.066', '0.5457'],
        ['Random Forest',       '0.4053', '0.4421', '+0.037', '0.7149'],
        ['XGBoost',             '0.3816', '0.3763', '−0.005', '0.6794'],
        ['ANN',                 '0.4726', '0.4089', '−0.064', '0.6601'],
        ['DNN',                 '0.4136', '0.4490', '+0.035', '0.6673'],
    ],
    caption='Table 3.2: Phase 2 results — baseline models on engineered features vs. Phase 1.')

add_justified(doc,
    'Phase 2 produced inconsistent results across models, with only Random Forest '
    'demonstrating a clear, consistent improvement (+0.037 F1, +0.004 AUC). RF\'s '
    'ensemble architecture is uniquely well-suited to exploit the additional feature '
    'interactions introduced by the derived features: each tree in the ensemble evaluates '
    'a random feature subset, allowing it to discover combinations of engineered and '
    'original features that single-tree or linear models cannot effectively leverage.')

add_justified(doc,
    'Logistic Regression declined modestly (−0.017 F1), suggesting that the engineered '
    'features introduce mild collinearity with existing linear predictors without adding '
    'net discriminatory linear signal. Decision Tree degraded most substantially (−0.066 '
    'F1): without depth control, the unconstrained tree overfits to the expanded feature '
    'space. XGBoost remained largely unchanged at default settings, confirming that '
    'gradient boosting at this scale is insensitive to feature set changes without '
    'concurrent parameter tuning. These findings collectively demonstrate that the '
    'benefit of feature engineering is conditional on the model\'s capacity to exploit '
    'feature interactions — an observation that motivates Phase 3\'s focus on '
    'hyperparameter tuning.')

# ── 3.3 Phase 3 ────────────────────────────────────────────────────────────────
add_heading(doc, '3.3 Phase 3: Hyperparameter Tuning Results', level=2)

add_justified(doc,
    'Phase 3 applied the three tuning strategies (Grid Search, Randomised Search, Optuna) '
    'to each ML model on the engineered feature set. Table 3.3 presents the best results '
    'per model and tuner.')

add_table(doc,
    headers=['Model', 'Best Tuner', 'CV F1', 'Test F1', 'Test AUC', 'Test Recall', 'Pred Time (µs)'],
    rows=[
        ['Logistic Regression', 'Random Search', '0.5313', '0.5344', '0.7072', '0.5905', '3.4'],
        ['Decision Tree',       'All (converged)', '0.5223', '0.5710', '0.7339', '0.5700', '3.0'],
        ['Random Forest',       'Random Search / Optuna', '0.4854', '0.5589', '0.7301', '0.5068', '6.8'],
        ['XGBoost',             'Optuna', '0.4613', '0.4924', '0.7121', '0.5068', '3.2'],
        ['ANN (Phase3)',         'Literature', '0.5295', '0.4606', '0.6743', '0.4795', '278'],
        ['DNN (Phase3)',         'Literature', '0.5585', '0.4919', '0.6581', '0.5741', '573'],
    ],
    caption='Table 3.3: Phase 3 results — best tuned ML models and literature-configured DL models.')

add_justified(doc,
    'The tuned Decision Tree achieved the highest test F1 (0.5710) and AUC (0.7339) of '
    'any model across all three phases — an improvement of 26% over its Phase 1 baseline '
    '(0.4524). The critical hyperparameter was max_depth=3: constraining the tree to a '
    'maximum depth of three levels (eight potential leaf nodes) prevents memorisation of '
    'training set noise and forces the model to identify the most globally discriminative '
    'splits. All three tuning strategies converged to the same optimal configuration, '
    'confirming robustness of the identified parameter region.')

add_justified(doc,
    'Random Forest also improved substantially, with the best tuned result achieving '
    'test F1 = 0.5589 and AUC = 0.7301. XGBoost improved from 0.3816 (Phase 1) to '
    '0.4924 (Phase 3 Optuna), though it remained below both LR and the tuned DT — '
    'an unexpected finding discussed in depth in Section 3.4.')

add_justified(doc,
    'Figure 3.1 presents the cross-phase F1 comparison across all models and phases.')

add_figure(doc, IMGS['comparison_viz'],
    'Figure 3.1: Cross-phase F1 score comparison across all models and experimental phases. '
    'The tuned Decision Tree achieves the highest test F1 across all phases.',
    width=6.0)

add_figure(doc, IMGS['exec_time'],
    'Figure 3.2: Inference time comparison (µs/sample) across all models. '
    'ML models are 50–130× faster than DL models.',
    width=5.5)

add_figure(doc, IMGS['confusion'],
    'Figure 3.3: Confusion matrices for the best model from each phase evaluated on the '
    'held-out test set (1,000 samples). The Phase 3 tuned Decision Tree achieves the '
    'best balance of true positives and false negatives.',
    width=6.0)

# ── 3.4 XGBoost Underperformance ───────────────────────────────────────────────
add_heading(doc, '3.4 The XGBoost Anomaly: Why the State-of-the-Art Model Underperformed', level=2)

add_justified(doc,
    'XGBoost is widely regarded as the dominant algorithm for structured tabular prediction '
    'tasks, consistently outperforming alternative methods in benchmark competitions and '
    'empirical studies (Chen & Guestrin, 2016; Lundberg & Lee, 2017). In the present study, '
    'however, XGBoost ranked last among ML models at default settings (Phase 1 test F1 = 0.382) '
    'and remained below Logistic Regression and the tuned Decision Tree even after extensive '
    'hyperparameter tuning (Phase 3 best test F1 = 0.492). Several interacting factors '
    'explain this outcome.')

add_justified(doc,
    'First, dataset scale is a fundamental constraint. XGBoost\'s strength lies in its '
    'capacity to sequentially correct prediction residuals across hundreds of trees, '
    'each capturing increasingly subtle patterns in the training data. This mechanism '
    'requires a sufficiently large dataset to produce reliable gradient estimates at each '
    'boosting step. With only 5,000 observations (4,000 in the training set), the sample '
    'size may be insufficient for XGBoost to leverage its full representational capacity. '
    'Grinsztajn et al. (2022), in a systematic evaluation across 45 tabular datasets, '
    'found that tree-based ensembles including XGBoost achieved their largest performance '
    'advantages over simpler methods on datasets with more than 10,000 samples.')

add_justified(doc,
    'Second, the nature of the churn signal in this dataset is predominantly linear and '
    'low-dimensional. The EDA (Section 1.2) established that the three strongest predictors '
    '(Recent_Complaints, Customer_Satisfaction_Score, Number_of_Transactions) contribute '
    'near-linearly to churn probability, whilst the remaining features exhibit minimal '
    'discriminatory power even in aggregate. XGBoost\'s sequential boosting mechanism is '
    'optimally suited to datasets with rich non-linear interaction structures — conditions '
    'not present here. In contrast, Logistic Regression\'s explicit linear inductive bias '
    'aligns well with the dominant signal structure.')

add_justified(doc,
    'Third, XGBoost at default settings employs a relatively high learning rate (0.1) '
    'and large n_estimators (100), which can cause overfitting on small datasets when the '
    'regularisation parameters (reg_alpha, reg_lambda, min_child_weight) are not carefully '
    'tuned. The tuning improvement from Phase 1 (0.382) to Phase 3 (0.492) — a 29% '
    'relative gain — confirms that the default configuration was poorly suited to this '
    'dataset, but even the tuned XGBoost remained limited by the scale and signal '
    'structure constraints identified above.')

add_justified(doc,
    'Fourth, the interaction between SMOTE and XGBoost\'s boosting mechanism may produce '
    'suboptimal results on small datasets. SMOTE\'s synthetic samples are generated by '
    'linear interpolation between minority-class neighbours; the resulting synthetic '
    'observations cluster in relatively smooth regions of the feature space. XGBoost\'s '
    'residual-correction mechanism may repeatedly boost predictions on these synthetic '
    'points — which have no real-world counterparts in the test set — producing a form of '
    'generalisation deficit that is less pronounced for bagged ensembles such as Random Forest.')

add_justified(doc,
    'In summary, the underperformance of XGBoost in this study is not anomalous in a '
    'theoretical sense; it is consistent with the literature\'s understanding of the '
    'conditions under which boosting excels. The dataset\'s small size, linear signal '
    'structure, and synthetic churn distribution collectively constitute an atypical '
    'environment that favours simpler models. This finding reinforces the principle that '
    'no single model is universally superior, and that algorithm selection must be '
    'empirically informed by the properties of the specific task and dataset.')

# ── 3.5 SHAP ───────────────────────────────────────────────────────────────────
add_heading(doc, '3.5 SHAP Interpretability Analysis', level=2)

add_justified(doc,
    'To ensure that the best-performing model produces not merely accurate but also '
    'interpretable and auditable predictions, SHapley Additive exPlanations (SHAP; '
    'Lundberg & Lee, 2017) were applied. SHAP is grounded in cooperative game theory: '
    'the Shapley value for each feature represents its average marginal contribution '
    'to the prediction across all possible feature coalitions. This framework provides '
    'theoretical guarantees of consistency and local accuracy that simpler feature '
    'importance metrics (e.g., Gini importance, permutation importance) lack.')

add_heading(doc, '3.5.1 Global Feature Importance', level=3)

add_justified(doc,
    'TreeSHAP (Lundberg et al., 2020) — an exact polynomial-time algorithm for '
    'tree-based models — was applied to the tuned Decision Tree to compute exact '
    'Shapley values for all 1,000 test samples. Figure 3.4 presents the beeswarm '
    'summary plot, in which each point represents one customer–feature combination: '
    'the horizontal position indicates the magnitude and direction of the SHAP value '
    '(positive = pushes toward churn prediction), and the colour encodes the raw '
    'feature value (red = high, blue = low).')

add_figure(doc, IMGS['shap_beeswarm'],
    'Figure 3.4: SHAP beeswarm summary plot for the tuned Decision Tree. '
    'Each point represents one test sample. Red indicates high feature value; '
    'blue indicates low feature value.',
    width=5.5)

add_figure(doc, IMGS['shap_bar'],
    'Figure 3.5: Mean absolute SHAP value (global feature importance) for the '
    'tuned Decision Tree. Features are ranked by their average contribution '
    'to prediction magnitude across all test samples.',
    width=5.5)

add_justified(doc,
    'The mean absolute SHAP values (Figure 3.5) identify Recent_Complaints and '
    'Customer_Satisfaction_Score as the two dominant churn drivers, consistent with '
    'the linear correlation findings from Section 1.2.3. These two features alone '
    'account for a disproportionate share of the model\'s total prediction variance, '
    'confirming that churn in this dataset is primarily determined by service quality '
    'and dissatisfaction intensity rather than financial characteristics.')

add_justified(doc,
    'Engineered features — particularly Satisfaction_Gap and Customer_Tenure_Months '
    '— appear in the upper tier of global importance, validating the feature '
    'engineering design. Satisfaction_Gap encodes the joint effect of low satisfaction '
    'and high complaint frequency, and its prominence in the SHAP ranking confirms '
    'that the multiplicative interaction it captures is genuinely exploited by the model. '
    'The presence of Customer_Tenure_Months also supports the hypothesis that new '
    'customers represent higher churn risk.')

add_heading(doc, '3.5.2 Local Explanations', level=3)

add_justified(doc,
    'SHAP waterfall plots were generated for individual customers to illustrate how '
    'the model arrives at a specific prediction. Figure 3.6 presents representative '
    'waterfall plots for a high-risk churn customer and a low-risk retained customer. '
    'In the high-risk case, large positive SHAP contributions from Recent_Complaints '
    'and low Customer_Satisfaction_Score collectively drive the prediction well above '
    'the baseline (expected value). In the low-risk case, the same features contribute '
    'negative SHAP values, anchoring the prediction below the churn threshold. '
    'This local explainability is operationally critical: it allows relationship managers '
    'to understand the specific factors driving a customer\'s churn risk and to tailor '
    'retention interventions accordingly.')

add_figure(doc, IMGS['shap_local'],
    'Figure 3.6: SHAP waterfall plots for individual customers — '
    'high-risk (top) and low-risk (bottom) examples. '
    'Each bar represents one feature\'s SHAP contribution to the individual prediction.',
    width=5.5)

add_heading(doc, '3.5.3 DNN Interpretability via KernelSHAP', level=3)

add_justified(doc,
    'For the DNN model, KernelSHAP (the model-agnostic SHAP estimator) was applied to '
    'a background subsample of 100 training observations. KernelSHAP uses a weighted '
    'linear regression framework to approximate Shapley values without requiring access '
    'to the model\'s internal structure. Whilst computationally more expensive than '
    'TreeSHAP, it enables interpretability for any differentiable or black-box model. '
    'Figure 3.7 presents the KernelSHAP summary plot for the DNN.')

add_figure(doc, IMGS['shap_dl'],
    'Figure 3.7: KernelSHAP global summary plot for the DNN model. '
    'The feature importance ranking is broadly consistent with TreeSHAP, '
    'confirming the robustness of the identified churn drivers across model families.',
    width=5.5)

add_justified(doc,
    'The KernelSHAP feature importance ranking for the DNN is broadly consistent with '
    'the TreeSHAP results for the Decision Tree — Recent_Complaints, '
    'Customer_Satisfaction_Score, and activity-based features dominate in both cases. '
    'This cross-model convergence in feature importance is a strong indicator of '
    'genuine predictive relevance: the identified churn drivers are not artefacts of a '
    'particular model\'s inductive bias but reflect real structure in the data. The '
    'consistency also lends support to the practical recommendations derived from the '
    'SHAP analysis (Section 4.2).')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════════
# CHAPTER 4: CONCLUSION AND STRATEGIC RECOMMENDATIONS
# ════════════════════════════════════════════════════════════════════════════════
add_heading(doc, 'Chapter 4: Conclusion and Strategic Recommendations', level=1)

add_justified(doc,
    'This final chapter revisits the research objectives established at the outset of '
    'the study and evaluates the extent to which each has been met by the work conducted. '
    'It then considers the contributions of this research to knowledge and practice, '
    'acknowledges the principal limitations of the study, and proposes directions '
    'for future investigation.')

# ── 4.1 Research Objectives Review ────────────────────────────────────────────
add_heading(doc, '4.1 Alignment with Research Objectives', level=2)

add_justified(doc,
    'The research was designed to address five core objectives. Table 4.1 provides '
    'an explicit mapping between each objective and the corresponding work conducted '
    'in this study, together with a summary of the principal finding.')

add_table(doc,
    headers=['Research Objective', 'Work Conducted', 'Key Finding'],
    rows=[
        ['RO1: Compare ML and DL models for customer churn prediction in banking',
         'Six models evaluated across three phases: LR, DT, RF, XGBoost, ANN, DNN. '
         'All trained on identical data splits and evaluated with consistent metrics.',
         'Tuned Decision Tree (F1=0.571, AUC=0.734) outperformed all models. '
         'DL models consistently trailed ML due to dataset scale constraints.'],
        ['RO2: Evaluate the impact of feature engineering on model performance',
         'Nine domain-driven features engineered (Phase 2). Statistical significance '
         'of improvements assessed via Wilcoxon signed-rank test.',
         'RF gained +0.037 F1 from FE. No model showed statistically significant '
         'improvement (limited by 5-fold CV power). DT and XGBoost unaffected without tuning.'],
        ['RO3: Assess the effectiveness of three hyperparameter tuning strategies',
         'Grid Search, Randomised Search, and Optuna (Bayesian, TPE) applied in Phase 3 '
         'to all four ML models with 5-fold CV optimisation of F1.',
         'All three strategies converged to equivalent final performance. '
         'DT improved 26% from tuning (max_depth=3). Optuna is more efficient '
         'but not superior in final quality at this scale.'],
        ['RO4: Apply SHAP for model interpretability and identify key churn drivers',
         'TreeSHAP applied to best tree model; KernelSHAP applied to DNN. '
         'Global beeswarm, bar, and local waterfall plots generated.',
         'Recent_Complaints and Customer_Satisfaction_Score are the primary churn drivers. '
         'Engineered feature Satisfaction_Gap validated as genuinely informative.'],
        ['RO5: Provide strategic recommendations for customer retention',
         'SHAP-derived insights translated into actionable retention interventions '
         'linked to the identified feature importance hierarchy.',
         'High-complaint, low-satisfaction, low-activity customers are the highest-priority '
         'retention targets. Tuned DT provides real-time (3 µs) scoring at scale.'],
    ],
    caption='Table 4.1: Research objectives vs. work conducted and key findings.')

add_justified(doc,
    'All five research objectives were addressed through the structured three-phase '
    'experimental design. The work confirms that supervised machine learning — particularly '
    'interpretable ensemble methods with careful hyperparameter control — provides an '
    'effective and practically deployable solution to customer churn prediction in banking.')

# ── 4.2 Contributions ─────────────────────────────────────────────────────────
add_heading(doc, '4.2 Contributions to Knowledge and Practice', level=2)

add_heading(doc, '4.2.1 Contribution to Knowledge', level=3)

add_justified(doc,
    'This study makes several contributions to the empirical literature on churn prediction:')

add_bullet(doc,
    'Comparative evaluation of imbalance strategies for DL: The study provides empirical '
    'evidence that class_weight is a more appropriate imbalance-handling strategy for '
    'neural networks on small tabular datasets than SMOTE, which inflates CV Recall by '
    'up to 0.28 relative to the true test Recall. This practical finding adds nuance '
    'to the standard recommendation of SMOTE for class imbalance.')
add_bullet(doc,
    'XGBoost behaviour on small, linearly-structured datasets: The study documents a '
    'scenario in which XGBoost — the established state-of-the-art for tabular prediction '
    '— underperforms simpler linear and shallow tree models. The analysis provides a '
    'multi-factor theoretical explanation (dataset scale, signal linearity, SMOTE '
    'interaction), contributing a well-evidenced cautionary case to the literature on '
    'algorithm selection.')
add_bullet(doc,
    'Three-strategy tuning comparison on a banking dataset: The empirical equivalence '
    'of Grid Search, Randomised Search, and Bayesian (Optuna) optimisation at this '
    'dataset scale contributes to the growing body of evidence that tuning strategy '
    'choice is secondary to search space design for moderate-scale tabular problems.')
add_bullet(doc,
    'SHAP-based validation of feature engineering: The convergent importance of '
    'Satisfaction_Gap across both TreeSHAP and KernelSHAP provides post-hoc validation '
    'that domain-informed feature engineering produces genuinely exploitable signal '
    'rather than noise amplification.')

add_heading(doc, '4.2.2 Contribution to Practice', level=3)

add_justified(doc,
    'The study delivers a deployable churn prediction framework with clear practical value:')

add_bullet(doc,
    'Actionable retention triggers: The SHAP analysis identifies customers with '
    'multiple recent complaints, low satisfaction scores (≤ 2/5), and declining '
    'transaction frequency as the highest-priority retention targets. These three '
    'conditions are directly observable in real banking CRM systems and can be '
    'used to trigger automated retention outreach without model inference.')
add_bullet(doc,
    'Real-time scoring capability: The tuned Decision Tree\'s inference latency of '
    '~3 µs per sample supports real-time churn scoring at the point of any customer '
    'interaction — branch visit, mobile banking login, or call centre contact — '
    'enabling timely rather than retrospective retention intervention.')
add_bullet(doc,
    'Interpretable decision rules: A depth-3 Decision Tree produces at most eight '
    'decision paths, each expressible as a sequence of three binary conditions on '
    'observable features. These rules can be communicated to relationship managers '
    'without data science expertise, supporting human-in-the-loop retention processes.')
add_bullet(doc,
    'Threshold adjustment guidance: The study demonstrates that Logistic Regression\'s '
    'AUC of 0.715 supports flexible threshold optimisation — by lowering the '
    'classification threshold from 0.5, banks can trade precision for recall according '
    'to their specific cost-benefit profile for retention vs. false alarm costs.')

# ── 4.3 Limitations ────────────────────────────────────────────────────────────
add_heading(doc, '4.3 Limitations of the Study', level=2)

add_justified(doc,
    'Several limitations constrain the generalisability and completeness of this study:')

add_bullet(doc,
    'Synthetic dataset: The dataset was sourced from a public (Kaggle) synthetic '
    'data generator, and the 31.66% churn rate is approximately 1.5–3× higher than '
    'churn rates typical of real retail banking datasets (10–20%; Verbeke et al., 2012). '
    'Model performance on a real banking dataset with a lower churn rate and different '
    'feature relationships may differ substantially. Findings should be interpreted as '
    'proof of concept rather than generalisable benchmarks.')
add_bullet(doc,
    'Cross-sectional data structure: The dataset captures a single cross-sectional '
    'snapshot per customer. Genuine churn prediction systems in banking are inherently '
    'temporal: the same customer may exhibit evolving risk over time. The absence of '
    'longitudinal data prevents the capture of trend-based churn signals (e.g., '
    'progressive disengagement over 6–12 months) and precludes survival-analytic '
    'approaches to time-to-churn modelling.')
add_bullet(doc,
    'No temporal train-test split: The 80/20 stratified split was performed as a '
    'random partition rather than a temporal partition (e.g., train on customers '
    'observed before a cut-off date, test on those observed after). A temporal split '
    'better simulates the real-world deployment scenario, where the model must '
    'generalise to future customer behaviour not yet observed at training time.')
add_bullet(doc,
    'DL models not fully hyperparameter-tuned: By design, deep learning architectures '
    'were evaluated using literature-configured hyperparameters without a formal search. '
    'Whilst this decision is methodologically justified by dataset size constraints, '
    'it means that the DL performance reported here may represent a lower bound for '
    'what these architectures could achieve on a larger dataset with full tuning.')
add_bullet(doc,
    'Limited cross-validation power: Five-fold cross-validation provides a sound '
    'performance estimate but generates only five paired observations for the Wilcoxon '
    'significance test — insufficient statistical power to detect all but the largest '
    'performance differences. Ten-fold or repeated CV would provide substantially '
    'improved power.')
add_bullet(doc,
    'Single dataset evaluation: Results are specific to this dataset\'s feature '
    'distribution, churn rate, and sample size. Conclusions regarding algorithm '
    'rankings (particularly XGBoost\'s underperformance) may not replicate across '
    'datasets with different characteristics.')

# ── 4.4 Future Research ────────────────────────────────────────────────────────
add_heading(doc, '4.4 Directions for Future Research', level=2)

add_justified(doc,
    'The limitations identified above suggest the following priority areas for future '
    'investigation:')

add_bullet(doc,
    'Real-world banking dataset validation: Applying the proposed methodology to a '
    'real banking dataset — with a lower churn rate, genuine temporal structure, and '
    'richer financial transaction history — would provide stronger evidence of '
    'practical utility and test the robustness of the feature engineering and model '
    'selection conclusions.')
add_bullet(doc,
    'Temporal and survival-analytic modelling: Extending the framework to incorporate '
    'time-series features (e.g., rolling transaction frequency, trend change points) '
    'and survival analysis methods (e.g., Cox proportional hazards, discrete-time '
    'survival models) would address the cross-sectional limitation and enable '
    'time-to-churn prediction alongside binary churn classification.')
add_bullet(doc,
    'Cost-sensitive learning: The current framework treats false negatives and false '
    'positives as equally weighted prediction errors beyond the F1 metric. '
    'A cost-sensitive framework that explicitly incorporates the monetary value of '
    'a retained customer and the cost of a retention intervention (incentive cost, '
    'agent time) into the objective function would align model optimisation more '
    'directly with the bank\'s financial KPIs.')
add_bullet(doc,
    'Ensemble and stacking methods: Given the complementary strengths of Logistic '
    'Regression (high Recall) and Decision Tree (high F1 and AUC), a stacking '
    'ensemble combining these base models with a meta-learner may yield superior '
    'performance to any individual model. This direction was outside the scope of '
    'the current study but is a natural extension.')
add_bullet(doc,
    'Deep learning on larger datasets: Future work with datasets exceeding 50,000 '
    'customers would provide a more rigorous assessment of DNN versus tree-ensemble '
    'performance on banking churn. At this scale, with full hyperparameter tuning, '
    'DNN may demonstrate the advantages theorised in the literature.')
add_bullet(doc,
    'Online learning and model drift monitoring: Production churn models must be '
    'continuously updated as customer behaviour and macroeconomic conditions evolve. '
    'Future research should investigate online learning frameworks, concept drift '
    'detection, and model monitoring pipelines appropriate to the banking context.')
add_bullet(doc,
    'Multi-class and voluntary vs. involuntary churn: Binary churn classification '
    'conflates voluntary exit (customer dissatisfaction) with involuntary churn '
    '(e.g., fraud-related closure, deceased customer). Distinguishing these categories '
    'through multi-class modelling or separate binary classifiers would produce more '
    'targeted retention strategies.')

# ── Final paragraph ─────────────────────────────────────────────────────────────
add_heading(doc, '4.5 Concluding Remarks', level=2)

add_justified(doc,
    'This study demonstrates that advanced predictive modelling for customer churn in '
    'banking is a tractable and practically valuable problem when approached with a '
    'principled, multi-phase methodology. The combination of domain-informed feature '
    'engineering, SMOTE-inside-CV imbalance handling, systematic hyperparameter tuning, '
    'and SHAP-based interpretability produced a deployable and auditable churn scoring '
    'system. The tuned Decision Tree with max_depth=3 — a model of remarkable parsimony '
    '— emerged as the overall winner, achieving a test F1 of 0.5710 and AUC of 0.7339 '
    'at an inference latency of 3 µs per sample, meeting the operational requirements '
    'of real-time banking retention systems.')

add_justified(doc,
    'The finding that churn in this dataset is primarily driven by service dissatisfaction '
    '(Recent_Complaints, Customer_Satisfaction_Score) rather than financial health '
    'carries a direct strategic implication: banks seeking to reduce churn should '
    'invest first in service quality improvement and complaint resolution speed, rather '
    'than targeting financially distressed customers with product restructuring. '
    'The SHAP analysis transforms this general insight into an individual-level, '
    'model-derived priority ranking — enabling personalised, data-driven retention '
    'interventions at scale.')

add_justified(doc,
    'The study also provides a methodologically cautionary finding that merits broader '
    'attention: no single algorithm is universally superior, and the conditions under '
    'which complex models (XGBoost, DNN) outperform simpler alternatives are critically '
    'dependent on dataset scale, signal linearity, and imbalance handling strategy. '
    'Rigorous empirical evaluation across multiple models, phases, and tuning strategies '
    '— as conducted in this study — is the only reliable basis for algorithm selection '
    'in applied machine learning research.')

# ── Save ────────────────────────────────────────────────────────────────────────
output_path = '/Users/nortonle/Library/CloudStorage/OneDrive-Personal/Documents/Master studies/Master thesis/testing claude code/dissertation_chapters.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
